from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def build_guidelines_doc():
    doc = Document()

    title = doc.add_heading("Job Description Writing Guidelines", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Best Practices Based on NLP Analysis of 200 Job Descriptions").paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading("1. Structure & Length", 1)
    doc.add_paragraph("Every job description should follow this structure in order:")
    structure = ["1. Engaging introduction (2–3 sentences)", "2. Key responsibilities (4–6 bullet points)",
                 "3. Requirements (3–5 bullet points — only what is truly necessary)",
                 "4. What We Offer / Benefits (3–5 bullet points)"]
    for s in structure:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph("Target length: 300–500 words. JDs under 200 or over 700 words perform significantly worse.")

    doc.add_heading("2. Readability Rules", 1)
    readability_rules = [
        "Target a Flesch-Kincaid Grade Level of 10–12 (use the textstat library to check).",
        "Keep average sentence length under 20 words.",
        "Avoid complex jargon unless it is truly necessary for the role.",
        "Use active voice: 'You will lead...' instead of 'The candidate will be responsible for leading...'",
        "Break up long paragraphs — use bullet points for lists of 3 or more items.",
    ]
    for r in readability_rules:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_heading("3. Tone & Language", 1)
    doc.add_heading("3.1 Use inclusive, welcoming language", 2)
    inclusive_dos = [
        "Use 'we' and 'you' — address the candidate directly.",
        "Describe what the candidate will gain, not just what they must do.",
        "Use 'ideally' or 'nice to have' for non-essential requirements.",
        "Mention team culture, collaboration, and growth opportunities.",
    ]
    for d in inclusive_dos:
        doc.add_paragraph(d, style='List Bullet')

    doc.add_heading("3.2 Words and phrases to avoid", 2)
    avoid_words = [
        "'Rockstar', 'ninja', 'wizard' — these discourage many candidates from applying.",
        "'Must', 'mandatory', 'required' for non-essential items — softens unnecessarily high bar.",
        "'24/7 availability' — unless genuinely required, this deters quality candidates.",
        "'Aggressive' targets — replace with 'ambitious' or 'stretching'.",
        "Gendered language — review with a gender decoder tool before posting.",
        "Vague filler phrases: 'duties as assigned', 'results-driven', 'fast-paced environment'.",
    ]
    for a in avoid_words:
        doc.add_paragraph(a, style='List Bullet')

    doc.add_heading("4. Requirements Section", 1)
    doc.add_paragraph("This is the section most likely to reduce applications if written poorly.")
    req_rules = [
        "List only genuine requirements — each extra 'must have' reduces your applicant pool.",
        "Separate 'essential' from 'desirable' requirements clearly.",
        "Accept equivalent experience in place of formal qualifications where possible.",
        "Avoid specifying exact years of experience — focus on skill level instead.",
        "Do not list tools or technologies that can be learned on the job.",
    ]
    for r in req_rules:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_heading("5. Benefits Section", 1)
    doc.add_paragraph(
        "Analysis shows JDs with a clear benefits section receive ~15% more applications. "
        "Never leave this section blank or use vague phrases like 'competitive salary'."
    )
    benefit_tips = [
        "Always mention salary range or at least 'competitive salary + bonus'.",
        "Highlight flexibility (remote work, flexible hours) — this is now a top candidate priority.",
        "Mention learning & development opportunities — especially attractive to high performers.",
        "Include work culture indicators: team size, collaboration style, values.",
    ]
    for b in benefit_tips:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_heading("6. Before You Post — Checklist", 1)
    checklist = [
        "☐  Is the JD between 300–500 words?",
        "☐  Does it follow the 4-section structure?",
        "☐  Have you checked Flesch-Kincaid Grade is 10–12?",
        "☐  Is the tone welcoming and positive (VADER score > 0.1)?",
        "☐  Have you removed all 'rockstar/ninja/mandatory' language?",
        "☐  Is there a clear benefits section?",
        "☐  Have you run it through a gender decoder?",
        "☐  Have requirements been trimmed to only genuine needs?",
    ]
    for item in checklist:
        doc.add_paragraph(item)

    os.makedirs("reports", exist_ok=True)
    doc.save("reports/JD_Writing_Guidelines.docx")
    print("✅ Guidelines saved: reports/JD_Writing_Guidelines.docx")

if __name__ == "__main__":
    build_guidelines_doc()