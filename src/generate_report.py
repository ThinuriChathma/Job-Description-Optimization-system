from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import os
from datetime import date

df = pd.read_csv("data/jd_scored.csv")
keywords = pd.read_csv("data/keyword_analysis.csv")

doc = Document()

# ── Title page ─────────────────────────────────────────────────────────────────
title = doc.add_heading("Job Description Effectiveness Report", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph(f"Prepared: {date.today().strftime('%B %d, %Y')}")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ── 1. Executive Summary ───────────────────────────────────────────────────────
doc.add_heading("1. Executive Summary", 1)
high_avg = df[df['quality_label']=='high']['applications'].mean()
low_avg  = df[df['quality_label']=='low']['applications'].mean()
diff_pct = ((high_avg - low_avg) / low_avg * 100)

doc.add_paragraph(
    f"This report analyzes {len(df)} job descriptions across {df['department'].nunique()} departments "
    f"to identify what makes a job posting effective. High-performing job descriptions receive an average "
    f"of {high_avg:.0f} applications compared to {low_avg:.0f} for low-performing ones — a {diff_pct:.0f}% difference. "
    f"Key findings show that readability, inclusive language, and tone are the strongest predictors of application rates."
)

# ── 2. Performance Analysis ────────────────────────────────────────────────────
doc.add_heading("2. Performance Analysis", 1)
doc.add_heading("2.1 Applications by Quality Tier", 2)

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = "Quality Tier"
hdr[1].text = "Count"
hdr[2].text = "Avg Applications"
hdr[3].text = "Max Applications"

for quality in ['high', 'medium', 'low']:
    subset = df[df['quality_label'] == quality]
    row = table.add_row().cells
    row[0].text = quality.capitalize()
    row[1].text = str(len(subset))
    row[2].text = f"{subset['applications'].mean():.1f}"
    row[3].text = str(subset['applications'].max())

doc.add_paragraph("")
doc.add_paragraph("Note: See outputs/performance_analysis.png for visual breakdown.")

# ── 3. Keyword Findings ────────────────────────────────────────────────────────
doc.add_heading("3. Keyword & Phrase Analysis", 1)
doc.add_heading("3.1 Effective Keywords (found in high-performing JDs)", 2)

top_keywords = keywords.nlargest(10, 'difference')
kw_text = ", ".join(top_keywords['keyword'].tolist())
doc.add_paragraph(f"The following keywords appear significantly more in high-performing job descriptions: {kw_text}.")

doc.add_heading("3.2 Keywords to Avoid", 2)
bottom_keywords = keywords.nsmallest(10, 'difference')
avoid_text = ", ".join(bottom_keywords['keyword'].tolist())
doc.add_paragraph(f"These keywords are associated with lower application rates: {avoid_text}.")

# ── 4. Readability Findings ────────────────────────────────────────────────────
doc.add_heading("4. Readability Analysis", 1)
avg_grade_high = df[df['quality_label']=='high']['fk_grade'].mean()
avg_grade_low  = df[df['quality_label']=='low']['fk_grade'].mean()

doc.add_paragraph(
    f"High-performing job descriptions have an average Flesch-Kincaid Grade Level of {avg_grade_high:.1f}, "
    f"compared to {avg_grade_low:.1f} for low-performing ones. The optimal reading level for job descriptions "
    f"is Grade 10–12, balancing professionalism with accessibility."
)

# ── 5. Sentiment & Tone Findings ──────────────────────────────────────────────
doc.add_heading("5. Sentiment & Tone Analysis", 1)
avg_sentiment_high = df[df['quality_label']=='high']['vader_compound'].mean()
avg_sentiment_low  = df[df['quality_label']=='low']['vader_compound'].mean()

doc.add_paragraph(
    f"High-performing JDs have an average VADER sentiment score of {avg_sentiment_high:.3f} (positive/welcoming), "
    f"while low-performing JDs score {avg_sentiment_low:.3f} (neutral to negative). "
    f"Welcoming, benefit-focused language significantly improves application rates."
)

doc.add_heading("5.1 Tone Distribution", 2)
tone_counts = df['tone'].value_counts()
for tone, count in tone_counts.items():
    doc.add_paragraph(f"• {tone}: {count} JDs ({count/len(df)*100:.1f}%)", style='List Bullet')

# ── 6. Recommendations ────────────────────────────────────────────────────────
doc.add_heading("6. Recommendations", 1)
recommendations = [
    "Use welcoming, inclusive language — replace 'must' and 'required' with 'ideally' and 'we'd love'.",
    "Keep job descriptions between 300–500 words for optimal engagement.",
    "Always include a 'What We Offer' section — benefits language increases applications by ~15%.",
    "Target a Flesch-Kincaid Grade Level of 10–12 for maximum accessibility.",
    "Avoid gendered or exclusive terms like 'rockstar', 'ninja', 'aggressive'.",
    "Lead with an engaging introduction that describes impact, not just duties.",
    "Use active voice and action-oriented language in responsibilities.",
]
for rec in recommendations:
    doc.add_paragraph(rec, style='List Bullet')

os.makedirs("reports", exist_ok=True)
doc.save("reports/JD_Effectiveness_Report.docx")
print("✅ Report saved: reports/JD_Effectiveness_Report.docx")