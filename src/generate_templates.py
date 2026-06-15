from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

TEMPLATES = {
    "Technical Role (e.g. Software Engineer, DevOps)": {
        "intro": "We are looking for a talented {Job Title} to join our engineering team. You'll work on meaningful technical challenges alongside a collaborative group of engineers who care about quality, learning, and impact.",
        "responsibilities": [
            "Design, build, and maintain scalable and reliable systems.",
            "Collaborate with product and design teams to deliver great user experiences.",
            "Write clean, well-tested code and participate in code reviews.",
            "Contribute to architectural decisions and technical roadmaps.",
            "Mentor teammates and actively share knowledge.",
        ],
        "requirements": [
            "3+ years of relevant software engineering experience.",
            "Strong problem-solving skills and attention to detail.",
            "Experience with [relevant tech stack] — we'll consider equivalent skills.",
            "Clear communication and a team-first mindset.",
        ],
        "offer": [
            "Competitive salary with annual performance review.",
            "Flexible working hours and remote-friendly culture.",
            "Annual learning budget for courses, books, and conferences.",
            "Health, dental, and vision insurance.",
            "25 days paid leave plus public holidays.",
        ]
    },
    "Business / Managerial Role (e.g. Product Manager, Operations Manager)": {
        "intro": "We're hiring an experienced {Job Title} to help shape strategy, lead cross-functional initiatives, and drive meaningful outcomes for our customers and business.",
        "responsibilities": [
            "Define and communicate a clear vision for your area of ownership.",
            "Work closely with stakeholders across engineering, design, sales, and marketing.",
            "Analyze data and user feedback to prioritize and make informed decisions.",
            "Lead projects end-to-end and ensure timely, high-quality delivery.",
            "Build and develop a high-performing team through coaching and feedback.",
        ],
        "requirements": [
            "5+ years in a relevant role with demonstrated leadership experience.",
            "Strong analytical skills and comfort working with data.",
            "Excellent communication and stakeholder management skills.",
            "Experience working in a fast-paced, cross-functional environment.",
        ],
        "offer": [
            "Competitive salary + performance bonus.",
            "Flexible and hybrid work arrangements.",
            "Opportunity to shape strategy and grow with the company.",
            "Comprehensive benefits including health and wellness support.",
            "Collaborative, inclusive team culture.",
        ]
    },
    "Entry-Level / Graduate Role": {
        "intro": "This is a fantastic opportunity for a motivated {Job Title} to launch their career with us. You don't need years of experience — we're looking for curiosity, drive, and a willingness to learn.",
        "responsibilities": [
            "Support the team with day-to-day tasks and ongoing projects.",
            "Learn from experienced colleagues in a structured onboarding program.",
            "Contribute ideas and bring a fresh perspective to the team.",
            "Help prepare reports, presentations, and research as needed.",
            "Grow into increasing responsibility over time.",
        ],
        "requirements": [
            "Degree in a relevant field (or equivalent experience).",
            "Enthusiastic, proactive attitude and eagerness to learn.",
            "Good written and verbal communication skills.",
            "Ability to work as part of a team and manage your own time.",
        ],
        "offer": [
            "Competitive entry-level salary.",
            "Structured mentorship and career development programme.",
            "Regular feedback and performance reviews.",
            "Access to learning resources and training days.",
            "Friendly, supportive team environment.",
        ]
    },
    "Sales / Client-Facing Role": {
        "intro": "We're looking for an energetic {Job Title} who loves building relationships and helping customers find the right solutions. You'll be joining a supportive, target-driven team that celebrates wins together.",
        "responsibilities": [
            "Develop and manage relationships with prospective and existing clients.",
            "Understand client needs and present tailored solutions clearly.",
            "Meet and exceed monthly and quarterly sales targets.",
            "Maintain accurate records in our CRM system.",
            "Collaborate with marketing and product teams on campaigns and feedback.",
        ],
        "requirements": [
            "2+ years of sales or client-facing experience preferred.",
            "Excellent interpersonal and communication skills.",
            "Self-motivated with a positive, resilient attitude.",
            "Comfortable with CRM tools and data-driven selling.",
        ],
        "offer": [
            "Base salary + uncapped commission structure.",
            "Clear career progression pathway.",
            "Regular team incentives and recognition programme.",
            "Full training and ongoing coaching.",
            "Collaborative and energetic team culture.",
        ]
    },
}

def build_template_doc():
    doc = Document()

    title = doc.add_heading("Optimized Job Description Templates Library", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph(
        "This library provides ready-to-use job description templates for common role categories. "
        "Each template is based on analysis of high-performing JDs and follows proven best practices "
        "for readability, inclusive language, and candidate engagement. "
        "Replace the placeholder text in [brackets] with role-specific details."
    )

    doc.add_page_break()

    for i, (template_name, content) in enumerate(TEMPLATES.items(), 1):
        doc.add_heading(f"Template {i}: {template_name}", 1)

        doc.add_heading("Introduction", 2)
        doc.add_paragraph(content["intro"])

        doc.add_heading("Key Responsibilities", 2)
        for item in content["responsibilities"]:
            doc.add_paragraph(item, style='List Bullet')

        doc.add_heading("What We're Looking For", 2)
        for item in content["requirements"]:
            doc.add_paragraph(item, style='List Bullet')

        doc.add_heading("What We Offer", 2)
        for item in content["offer"]:
            doc.add_paragraph(item, style='List Bullet')

        if i < len(TEMPLATES):
            doc.add_page_break()

    os.makedirs("reports", exist_ok=True)
    doc.save("reports/JD_Templates_Library.docx")
    print("✅ Templates library saved: reports/JD_Templates_Library.docx")

if __name__ == "__main__":
    build_template_doc()